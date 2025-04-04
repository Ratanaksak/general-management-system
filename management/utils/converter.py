import os
from pdf2image import convert_from_path
from docx import Document
from PIL import Image
import pytesseract
from django.conf import settings

def convert_pdf_to_word(pdf_path):
    try:
        doc = Document()
        images = convert_from_path(
            pdf_path,
            poppler_path=settings.POPPLER_PATH  # Use the path from settings
        )
        
        for i, image in enumerate(images):
            temp_img_path = os.path.join(settings.MEDIA_ROOT, f'temp_page_{i}.jpg')
            image.save(temp_img_path, 'JPEG')
            
            try:
                text = pytesseract.image_to_string(Image.open(temp_img_path))
                doc.add_paragraph(text)
                if i < len(images) - 1:
                    doc.add_page_break()
            finally:
                if os.path.exists(temp_img_path):
                    os.remove(temp_img_path)
                    
        return doc
    except Exception as e:
        raise Exception(f"PDF conversion failed: {str(e)}")

def convert_image_to_word(image_path):
    try:
        doc = Document()
        text = pytesseract.image_to_string(Image.open(image_path))
        doc.add_paragraph(text)
        return doc
    except Exception as e:
        raise Exception(f"Image conversion failed: {str(e)}")